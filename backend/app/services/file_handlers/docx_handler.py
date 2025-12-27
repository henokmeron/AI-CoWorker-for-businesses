"""
DOCX handler using python-docx.
"""
import logging
from typing import Dict, Any, List
from pathlib import Path

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .base_handler import BaseFileHandler

logger = logging.getLogger(__name__)


class DOCXHandler(BaseFileHandler):
    """
    DOCX handler using python-docx.
    """
    
    SUPPORTED_TYPES = ['docx', 'doc']
    
    def __init__(self):
        """Initialize the DOCX handler."""
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx library not available. "
                "Install with: pip install python-docx"
            )
    
    def can_handle(self, file_path: str, file_type: str) -> bool:
        """Check if file type is supported."""
        return file_type.lower() in ['docx', 'doc'] and DOCX_AVAILABLE
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported file types."""
        return self.SUPPORTED_TYPES if DOCX_AVAILABLE else []
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """
        Process DOCX file using python-docx.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        try:
            logger.info(f"Processing DOCX with python-docx: {file_path}")
            
            doc = DocxDocument(file_path)
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            if table_text:
                text_content.extend(table_text)
            
            full_text = "\n\n".join(text_content)
            
            if not full_text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")
                full_text = "[DOCX file - text extraction may be limited]"
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "has_text": bool(full_text.strip())
            }
            
            # Extract document properties if available
            if doc.core_properties:
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
            
            # Create simple chunks
            chunks = self._create_simple_chunks(full_text)
            
            return {
                "text": full_text,
                "metadata": metadata,
                "chunks": chunks
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise RuntimeError(f"Failed to process DOCX: {str(e)}")
    
    def _create_simple_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """
        Create simple text chunks.
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of chunks with text and metadata
        """
        chunks = []
        start = 0
        text_length = len(text)
        chunk_index = 0
        
        while start < text_length:
            end = start + chunk_size
            chunk_text = text[start:end]
            
            # Try to break at sentence boundary
            if end < text_length:
                last_period = chunk_text.rfind('. ')
                last_newline = chunk_text.rfind('\n\n')
                break_point = max(last_period, last_newline)
                
                if break_point > chunk_size * 0.5:
                    chunk_text = chunk_text[:break_point + 1]
                    end = start + break_point + 1
            
            chunks.append({
                "text": chunk_text.strip(),
                "metadata": {
                    "chunk_index": chunk_index,
                    "start_char": start,
                    "end_char": end,
                }
            })
            
            chunk_index += 1
            start = end - overlap
        
        return chunks

