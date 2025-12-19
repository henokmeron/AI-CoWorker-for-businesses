"""
Script to create comprehensive Word document for AI CoWorker Platform Plan.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    heading.style.font.size = Pt(14 if level == 1 else 12)
    return heading

def add_bullet_point(doc, text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_numbered_point(doc, text, bold_prefix=None):
    """Add a numbered point."""
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def create_plan_document():
    """Create the comprehensive plan document."""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('AI CoWorker Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.bold = True
    
    subtitle = doc.add_heading('Comprehensive Implementation Plan & Feature Specification', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_paragraph()  # Spacing
    
    # Table of Contents placeholder
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        "1. Vision & Purpose",
        "2. Core Architecture",
        "3. Feature Specifications",
        "4. User Interface & Experience",
        "5. Authentication & Security",
        "6. Document Management",
        "7. GPT/Custom Assistant Management",
        "8. Conversation Management",
        "9. Settings & Configuration",
        "10. Integration Capabilities",
        "11. Advanced Features",
        "12. Implementation Roadmap",
        "13. Technical Requirements"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.style.font.size = Pt(11)
    
    doc.add_page_break()
    
    # 1. VISION & PURPOSE
    doc.add_heading('1. Vision & Purpose', 1)
    
    doc.add_paragraph(
        "The AI CoWorker Platform is a sophisticated, business-grade AI assistant system designed to "
        "replicate and enhance the ChatGPT experience for enterprise use. The platform serves as a "
        "persistent digital employee that learns immediately from uploaded documents, reasons across "
        "full context (not just keywords), and assists businesses at the director level."
    )
    
    doc.add_heading('1.1 Primary Purpose', 2)
    doc.add_paragraph(
        "To deliver the most capable AI coworker on the market: a persistent digital employee that "
        "learns immediately from uploaded documents, reasons across full context, and assists "
        "businesses with intelligent, context-aware responses grounded in their specific business knowledge."
    )
    
    doc.add_heading('1.2 Target Users', 2)
    add_bullet_point(doc, "Business directors and executives")
    add_bullet_point(doc, "Operations managers")
    add_bullet_point(doc, "Administrative staff")
    add_bullet_point(doc, "Teams requiring instant access to business documents and policies")
    
    doc.add_heading('1.3 Key Value Propositions', 2)
    add_bullet_point(doc, "ChatGPT-style interface with enterprise-grade features")
    add_bullet_point(doc, "Persistent document storage (never deleted when chat ends)")
    add_bullet_point(doc, "Custom GPT/Assistant configurations per business")
    add_bullet_point(doc, "Sophisticated reasoning across full document context")
    add_bullet_point(doc, "Integration with Microsoft Outlook and Gmail")
    add_bullet_point(doc, "Browser-based and add-in accessible")
    add_bullet_point(doc, "No placeholder features - everything works")
    
    doc.add_page_break()
    
    # 2. CORE ARCHITECTURE
    doc.add_heading('2. Core Architecture', 1)
    
    doc.add_heading('2.1 System Architecture', 2)
    doc.add_paragraph(
        "The platform uses an Agent + Memory + Tools architecture:"
    )
    add_bullet_point(doc, "LLM Layer: For reasoning and generation (OpenAI, Anthropic, or Ollama)")
    add_bullet_point(doc, "Vector Database: For semantic document memory (ChromaDB)")
    add_bullet_point(doc, "Structured Memory: For rules, tone, permissions, and business-specific configurations")
    add_bullet_point(doc, "Agent Loop: Retrieve → Reason → Act → Ground responses in business knowledge")
    
    doc.add_heading('2.2 Technology Stack', 2)
    add_bullet_point(doc, "Backend: FastAPI (Python) - RESTful API")
    add_bullet_point(doc, "Frontend: Streamlit (Python) - Web interface")
    add_bullet_point(doc, "Vector DB: ChromaDB - Document embeddings and semantic search")
    add_bullet_point(doc, "Document Processing: Unstructured.io, PyPDF2 - Multi-format support")
    add_bullet_point(doc, "LLM Integration: OpenAI GPT-4, Anthropic Claude, Ollama (local)")
    add_bullet_point(doc, "Authentication: JWT tokens, OAuth 2.0 (Google, Microsoft)")
    add_bullet_point(doc, "Storage: File system + JSON metadata (scalable to PostgreSQL)")
    
    doc.add_heading('2.3 Data Flow', 2)
    doc.add_paragraph(
        "1. User uploads document → 2. Document processed and chunked → 3. Embedded into vector DB → "
        "4. User asks question → 5. Query embedded → 6. Similar chunks retrieved → 7. LLM generates "
        "response grounded in retrieved context → 8. Response displayed with source citations"
    )
    
    doc.add_page_break()
    
    # 3. FEATURE SPECIFICATIONS
    doc.add_heading('3. Feature Specifications', 1)
    
    doc.add_heading('3.1 Core Features Overview', 2)
    add_bullet_point(doc, "My GPTs / Custom Assistants - Create and manage multiple AI assistants")
    add_bullet_point(doc, "Conversations - Persistent chat threads per GPT")
    add_bullet_point(doc, "Document Upload & Management - Upload, process, and attach documents to GPTs")
    add_bullet_point(doc, "Settings - Comprehensive configuration panel")
    add_bullet_point(doc, "Authentication - Login, logout, user management")
    add_bullet_point(doc, "Integration - Outlook add-in, Gmail extension, browser access")
    
    doc.add_page_break()
    
    # 4. USER INTERFACE & EXPERIENCE
    doc.add_heading('4. User Interface & Experience', 1)
    
    doc.add_heading('4.1 Layout Structure (ChatGPT-Style)', 2)
    doc.add_paragraph(
        "The interface must replicate ChatGPT's layout exactly:"
    )
    add_bullet_point(doc, "Left Sidebar: Contains 'My GPTs' section and 'Conversations' section")
    add_bullet_point(doc, "Main Chat Area: Central area for displaying messages and chat history")
    add_bullet_point(doc, "Prompt Field: Bottom of main area with '+' button to the left")
    add_bullet_point(doc, "Avatar Button: Fixed at bottom-left of sidebar (not main window)")
    
    doc.add_heading('4.2 Sidebar Components', 2)
    
    doc.add_heading('4.2.1 "My GPTs" Section', 3)
    doc.add_paragraph("Located at the top of the sidebar:")
    add_bullet_point(doc, "Header: 'My GPTs' with '+' button to create new GPT")
    add_bullet_point(doc, "GPT List: Each GPT displayed as a card/item")
    add_bullet_point(doc, "GPT Kebab Menu (⋮): Three-dot menu on each GPT with options:")
    doc.add_paragraph("  • New chat - Start new conversation with same GPT", style='List Bullet 2')
    doc.add_paragraph("  • About - Show GPT description and details", style='List Bullet 2')
    doc.add_paragraph("  • Edit GPT - Open edit panel for GPT configuration", style='List Bullet 2')
    doc.add_paragraph("  • Hide - Hide GPT from list (role-aware)", style='List Bullet 2')
    doc.add_paragraph("  • Delete - Delete GPT and all associated data", style='List Bullet 2')
    add_bullet_point(doc, "GPT Selection: Clicking GPT name selects it and loads its conversations")
    
    doc.add_heading('4.2.2 "Conversations" Section', 3)
    doc.add_paragraph("Located below 'My GPTs' in sidebar:")
    add_bullet_point(doc, "Header: 'Conversations' (or 'Chat History')")
    add_bullet_point(doc, "Conversation List: All conversations for selected GPT")
    add_bullet_point(doc, "Conversation Menu (⋯): Three-dot menu on each conversation with options:")
    doc.add_paragraph("  • Rename - Change conversation title", style='List Bullet 2')
    doc.add_paragraph("  • Archive - Move conversation to archive", style='List Bullet 2')
    doc.add_paragraph("  • Delete - Delete conversation (NOT the GPT)", style='List Bullet 2')
    add_bullet_point(doc, "Conversation Selection: Clicking conversation loads its chat history")
    
    doc.add_heading('4.2.3 Avatar Button (Bottom-Left of Sidebar)', 3)
    doc.add_paragraph("Fixed at the absolute bottom of the sidebar:")
    add_bullet_point(doc, "Display: Circular button showing user initials when logged in, '👤' icon when logged out")
    add_bullet_point(doc, "Styling: Green/primary color when logged in, gray/secondary when logged out")
    add_bullet_point(doc, "Click Action: Opens account dropdown menu in sidebar (not main window)")
    add_bullet_point(doc, "Dropdown Options (when logged in):")
    doc.add_paragraph("  • ⚙️ Settings - Open settings panel", style='List Bullet 2')
    doc.add_paragraph("  • ⬆️ Upgrade - Subscription/upgrade options", style='List Bullet 2')
    doc.add_paragraph("  • ❓ Help - Help center and documentation", style='List Bullet 2')
    doc.add_paragraph("  • 🚪 Log out - Sign out of account", style='List Bullet 2')
    add_bullet_point(doc, "Dropdown Options (when logged out):")
    doc.add_paragraph("  • Login - Show login form", style='List Bullet 2')
    doc.add_paragraph("  • Sign up - Show registration form", style='List Bullet 2')
    
    doc.add_heading('4.3 Main Chat Area', 2)
    add_bullet_point(doc, "Chat History: Displays all messages in current conversation")
    add_bullet_point(doc, "Message Format: User messages on right, AI responses on left (or vice versa)")
    add_bullet_point(doc, "Source Citations: AI responses show source documents when applicable")
    add_bullet_point(doc, "File Attachments: Display uploaded files in chat history")
    add_bullet_point(doc, "Empty State: Welcome message when no conversation selected")
    
    doc.add_heading('4.4 Prompt Field & File Upload', 2)
    doc.add_paragraph("Located at the bottom of main chat area:")
    add_bullet_point(doc, "'+' Button: Positioned to the LEFT of the prompt field (not top corner)")
    add_bullet_point(doc, "'+' Button Action: Toggles file uploader above prompt field")
    add_bullet_point(doc, "File Uploader: Appears ABOVE prompt when '+' is clicked (not in chat history)")
    add_bullet_point(doc, "Supported Formats: PDF, DOCX, TXT, XLSX, DOC, XLS, PPTX, CSV")
    add_bullet_point(doc, "File Size Limit: 200MB per file")
    add_bullet_point(doc, "Upload Process: Drag-and-drop or browse files")
    add_bullet_point(doc, "Processing Feedback: Spinner and success/error messages")
    add_bullet_point(doc, "Chat Input: Standard text input with send button (paper airplane icon)")
    
    doc.add_heading('4.5 "Reply as Me" Toggle', 2)
    doc.add_paragraph("Located at the bottom of sidebar (above avatar):")
    add_bullet_point(doc, "Function: Toggle between personalized replies and categorization mode")
    add_bullet_point(doc, "ON: AI responds in personalized tone based on business settings")
    add_bullet_point(doc, "OFF: AI categorizes and summarizes only (no personalized responses)")
    add_bullet_point(doc, "State: Persists per business/GPT configuration")
    
    doc.add_page_break()
    
    # 5. AUTHENTICATION & SECURITY
    doc.add_heading('5. Authentication & Security', 1)
    
    doc.add_heading('5.1 Authentication Flow', 2)
    add_bullet_point(doc, "Email/Password Login: Standard form with validation")
    add_bullet_point(doc, "Google OAuth: Full OAuth 2.0 integration (not placeholder)")
    add_bullet_point(doc, "Microsoft OAuth: Full OAuth 2.0 integration (not placeholder)")
    add_bullet_point(doc, "JWT Tokens: Secure token-based authentication")
    add_bullet_point(doc, "Session Management: Persistent sessions with timeout")
    
    doc.add_heading('5.2 User Management', 2)
    add_bullet_point(doc, "User Registration: Sign up with email verification")
    add_bullet_point(doc, "Profile Management: Update name, email, password")
    add_bullet_point(doc, "Password Reset: Forgot password flow with email reset")
    add_bullet_point(doc, "Account Deletion: User can delete account and all data")
    
    doc.add_heading('5.3 Security Features', 2)
    add_bullet_point(doc, "API Key Authentication: Backend requires API key for all requests")
    add_bullet_point(doc, "Data Isolation: Each business has isolated document storage")
    add_bullet_point(doc, "Role-Based Access: Different permissions for different user roles")
    add_bullet_point(doc, "2FA Support: Two-factor authentication option")
    add_bullet_point(doc, "Session Timeout: Automatic logout after inactivity")
    
    doc.add_page_break()
    
    # 6. DOCUMENT MANAGEMENT
    doc.add_heading('6. Document Management', 1)
    
    doc.add_heading('6.1 Document Upload', 2)
    add_bullet_point(doc, "Upload Methods: Drag-and-drop, file browser, or programmatic API")
    add_bullet_point(doc, "Supported Formats: PDF, DOCX, TXT, XLSX, DOC, XLS, PPTX, CSV, and more")
    add_bullet_point(doc, "File Size: Up to 200MB per file")
    add_bullet_point(doc, "Batch Upload: Multiple files at once")
    add_bullet_point(doc, "Upload Location: Can upload in chat (general) or in Edit GPT (GPT-specific)")
    
    doc.add_heading('6.2 Document Processing', 2)
    add_bullet_point(doc, "Text Extraction: Extract text from all supported formats")
    add_bullet_point(doc, "Chunking: Split documents into semantic chunks (configurable size)")
    add_bullet_point(doc, "Embedding: Generate vector embeddings for each chunk")
    add_bullet_point(doc, "Storage: Store chunks in vector database with metadata")
    add_bullet_point(doc, "Metadata: Track file name, upload date, size, GPT association")
    
    doc.add_heading('6.3 Document Association', 2)
    add_bullet_point(doc, "GPT-Specific: Documents uploaded in Edit GPT are attached to that GPT")
    add_bullet_point(doc, "General Upload: Documents uploaded in chat are available to all GPTs (or specific business)")
    add_bullet_point(doc, "Namespace Isolation: Each GPT has its own vector namespace")
    add_bullet_point(doc, "Persistent Storage: Documents remain even when conversations are deleted")
    
    doc.add_heading('6.4 Document Management UI', 2)
    add_bullet_point(doc, "Document List: View all documents in Edit GPT panel")
    add_bullet_point(doc, "Delete Documents: Remove documents from GPT knowledge base")
    add_bullet_point(doc, "Document Preview: View document metadata and processing status")
    add_bullet_point(doc, "Re-processing: Option to re-process documents if needed")
    
    doc.add_page_break()
    
    # 7. GPT/CUSTOM ASSISTANT MANAGEMENT
    doc.add_heading('7. GPT/Custom Assistant Management', 1)
    
    doc.add_heading('7.1 GPT Lifecycle', 2)
    doc.add_paragraph(
        "CRITICAL: GPTs and Conversations are SEPARATE entities with different lifecycles:"
    )
    add_bullet_point(doc, "GPT Creation: GPT is created with name, description, instructions")
    add_bullet_point(doc, "GPT Persistence: GPT remains in 'My GPTs' list until explicitly deleted")
    add_bullet_point(doc, "Conversation Deletion: Deleting a conversation does NOT delete the GPT")
    add_bullet_point(doc, "GPT Deletion: Deleting a GPT removes the GPT AND all its conversations")
    add_bullet_point(doc, "New Chat: Creates new conversation using same GPT configuration and vectors")
    
    doc.add_heading('7.2 GPT Configuration', 2)
    add_bullet_point(doc, "Name: Display name for the GPT")
    add_bullet_point(doc, "Description: Brief description of GPT's purpose")
    add_bullet_point(doc, "Instructions: System prompt/instructions for the GPT")
    add_bullet_point(doc, "Knowledge Base: Documents attached to this GPT")
    add_bullet_point(doc, "Tone Templates: Response style templates (per business)")
    add_bullet_point(doc, "Settings: GPT-specific settings (temperature, model, etc.)")
    
    doc.add_heading('7.3 Edit GPT Panel', 2)
    doc.add_paragraph("Opens when 'Edit GPT' is clicked from kebab menu:")
    add_bullet_point(doc, "Side Panel: Slides in from right (or full-screen modal)")
    add_bullet_point(doc, "Fields: Name, Description, Instructions (all editable)")
    add_bullet_point(doc, "Knowledge Base Section:")
    doc.add_paragraph("  • List of uploaded documents", style='List Bullet 2')
    doc.add_paragraph("  • Upload new document button", style='List Bullet 2')
    doc.add_paragraph("  • Delete document button for each document", style='List Bullet 2')
    add_bullet_point(doc, "Save Button: Saves changes and applies instantly")
    add_bullet_point(doc, "Cancel Button: Closes panel without saving")
    add_bullet_point(doc, "Real Backend: All changes saved to GPT config store (not fake)")
    
    doc.add_heading('7.4 GPT Vector Namespace', 2)
    add_bullet_point(doc, "Isolation: Each GPT has its own vector database namespace")
    add_bullet_point(doc, "Reusability: Same GPT vectors used across all conversations for that GPT")
    add_bullet_point(doc, "Persistence: Vectors remain even when conversations are deleted")
    add_bullet_point(doc, "Update: Adding documents to GPT updates its vector namespace")
    
    doc.add_page_break()
    
    # 8. CONVERSATION MANAGEMENT
    doc.add_heading('8. Conversation Management', 1)
    
    doc.add_heading('8.1 Conversation Lifecycle', 2)
    add_bullet_point(doc, "Creation: New conversation created when starting new chat with a GPT")
    add_bullet_point(doc, "Title: Auto-generated from first message or user-defined")
    add_bullet_point(doc, "Messages: All chat messages stored in conversation")
    add_bullet_point(doc, "Deletion: Deleting conversation removes chat history only (NOT the GPT)")
    add_bullet_point(doc, "Archiving: Conversations can be archived (hidden but not deleted)")
    
    doc.add_heading('8.2 Conversation Features', 2)
    add_bullet_point(doc, "Chat History: Full message history with timestamps")
    add_bullet_point(doc, "Rename: User can rename conversation")
    add_bullet_point(doc, "Export: Export conversation as text or PDF")
    add_bullet_point(doc, "Search: Search within conversation messages")
    add_bullet_point(doc, "Source Citations: Each AI response shows source documents")
    
    doc.add_heading('8.3 "New Chat" Functionality', 2)
    doc.add_paragraph("When 'New chat' is clicked from GPT kebab menu:")
    add_bullet_point(doc, "Creates new conversation with same GPT")
    add_bullet_point(doc, "Reuses GPT configuration (instructions, settings)")
    add_bullet_point(doc, "Reuses GPT vector namespace (all documents still available)")
    add_bullet_point(doc, "Starts with empty chat history")
    add_bullet_point(doc, "Previous conversations remain accessible")
    
    doc.add_page_break()
    
    # 9. SETTINGS & CONFIGURATION
    doc.add_heading('9. Settings & Configuration', 1)
    
    doc.add_heading('9.1 Settings Panel Structure', 2)
    doc.add_paragraph("Tabbed layout with 7 main sections:")
    
    doc.add_heading('9.1.1 General Tab', 3)
    add_bullet_point(doc, "Language: Select interface language")
    add_bullet_point(doc, "Theme: Light/Dark mode toggle")
    add_bullet_point(doc, "Font Size: Adjustable text size")
    add_bullet_point(doc, "Auto-save: Enable/disable auto-save")
    
    doc.add_heading('9.1.2 Notifications Tab', 3)
    add_bullet_point(doc, "Email Notifications: Enable/disable email alerts")
    add_bullet_point(doc, "Browser Notifications: Enable/disable browser push notifications")
    add_bullet_point(doc, "Desktop Notifications: Enable/disable desktop notifications")
    add_bullet_point(doc, "Notification Preferences: Customize what triggers notifications")
    
    doc.add_heading('9.1.3 Personalization Tab', 3)
    add_bullet_point(doc, "Custom Instructions: Default instructions for all GPTs")
    add_bullet_point(doc, "Response Style: Preferred response tone and format")
    add_bullet_point(doc, "Tone Templates: Save and reuse tone templates per business")
    add_bullet_point(doc, "Personality: Adjust AI personality traits")
    
    doc.add_heading('9.1.4 App Connectors Tab', 3)
    add_bullet_point(doc, "Gmail Integration: Connect Gmail account")
    add_bullet_point(doc, "Outlook Integration: Connect Outlook/Microsoft account")
    add_bullet_point(doc, "Slack Integration: Connect Slack workspace")
    add_bullet_point(doc, "API Keys: Manage API keys for integrations")
    
    doc.add_heading('9.1.5 Data Control Tab', 3)
    add_bullet_point(doc, "Export Data: Export all user data (conversations, documents, settings)")
    add_bullet_point(doc, "Delete Data: Delete specific data or all data")
    add_bullet_point(doc, "Data Retention: Configure data retention policies")
    add_bullet_point(doc, "Backup: Manual and automatic backup options")
    
    doc.add_heading('9.1.6 Security Tab', 3)
    add_bullet_point(doc, "Change Password: Update account password")
    add_bullet_point(doc, "Two-Factor Authentication: Enable/disable 2FA")
    add_bullet_point(doc, "Session Timeout: Configure automatic logout")
    add_bullet_point(doc, "Active Sessions: View and manage active sessions")
    add_bullet_point(doc, "API Keys: Manage API keys for programmatic access")
    
    doc.add_heading('9.1.7 Account Tab', 3)
    add_bullet_point(doc, "Name: Update display name")
    add_bullet_point(doc, "Email: Update email address")
    add_bullet_point(doc, "Subscription: View and manage subscription")
    add_bullet_point(doc, "Billing: View billing history and payment methods")
    add_bullet_point(doc, "Delete Account: Permanently delete account")
    
    doc.add_heading('9.2 Settings Features', 2)
    add_bullet_point(doc, "Lazy Loading: Each tab loads content only when opened")
    add_bullet_point(doc, "State Persistence: Settings saved per tab")
    add_bullet_point(doc, "Real-time Updates: Changes apply instantly")
    add_bullet_point(doc, "Back Navigation: '← Back to Chat' button to return to main interface")
    
    doc.add_page_break()
    
    # 10. INTEGRATION CAPABILITIES
    doc.add_heading('10. Integration Capabilities', 1)
    
    doc.add_heading('10.1 Outlook Add-in', 2)
    add_bullet_point(doc, "Installation: Official Microsoft Office add-in")
    add_bullet_point(doc, "In-Client Panel: Embedded chat panel within Outlook")
    add_bullet_point(doc, "Email Analysis: Analyze email content and attachments")
    add_bullet_point(doc, "Document Processing: Upload and process attachments directly from Outlook")
    add_bullet_point(doc, "Context Awareness: AI understands email context and business documents")
    add_bullet_point(doc, "Quick Actions: Reply suggestions, summarization, categorization")
    
    doc.add_heading('10.2 Gmail Extension', 2)
    add_bullet_point(doc, "Browser Extension: Chrome/Edge extension for Gmail")
    add_bullet_point(doc, "Sidebar Integration: AI panel in Gmail sidebar")
    add_bullet_point(doc, "Email Analysis: Analyze emails and attachments")
    add_bullet_point(doc, "Document Processing: Process Gmail attachments")
    add_bullet_point(doc, "Quick Actions: Similar to Outlook add-in")
    
    doc.add_heading('10.3 Browser Access', 2)
    add_bullet_point(doc, "Web Application: Full-featured web app accessible via browser")
    add_bullet_point(doc, "Responsive Design: Works on desktop, tablet, and mobile")
    add_bullet_point(doc, "PWA Support: Progressive Web App for offline capability")
    add_bullet_point(doc, "Cross-Browser: Works on Chrome, Firefox, Safari, Edge")
    
    doc.add_heading('10.4 API Access', 2)
    add_bullet_point(doc, "REST API: Full RESTful API for programmatic access")
    add_bullet_point(doc, "Webhooks: Event-driven integrations")
    add_bullet_point(doc, "SDK: Software development kits for popular languages")
    add_bullet_point(doc, "Documentation: Comprehensive API documentation")
    
    doc.add_page_break()
    
    # 11. ADVANCED FEATURES
    doc.add_heading('11. Advanced Features', 1)
    
    doc.add_heading('11.1 Sophisticated Reasoning', 2)
    doc.add_paragraph(
        "The AI must provide sophisticated reasoning, not just keyword matching:"
    )
    add_bullet_point(doc, "Full Context Understanding: AI reads and understands entire documents")
    add_bullet_point(doc, "Cross-Document Reasoning: Connect information across multiple documents")
    add_bullet_point(doc, "Numerical Analysis: Understand and calculate from spreadsheets")
    add_bullet_point(doc, "Temporal Reasoning: Understand dates, schedules, and time-based information")
    add_bullet_point(doc, "Causal Reasoning: Understand cause-and-effect relationships")
    add_bullet_point(doc, "Multi-Step Reasoning: Break down complex questions into steps")
    
    doc.add_heading('11.2 RAG (Retrieval Augmented Generation) Flow', 2)
    doc.add_paragraph("Enforced retrieval-first approach:")
    add_bullet_point(doc, "Query Processing: User query is embedded")
    add_bullet_point(doc, "Semantic Search: Search vector database for relevant chunks")
    add_bullet_point(doc, "Context Retrieval: Retrieve top-k most relevant document chunks")
    add_bullet_point(doc, "LLM Generation: LLM generates response grounded in retrieved context")
    add_bullet_point(doc, "Source Attribution: Every response cites source documents")
    add_bullet_point(doc, "No Hallucination: AI cannot make up information not in documents")
    
    doc.add_heading('11.3 Response Modes', 2)
    doc.add_heading('11.3.1 "Reply as Me" Mode (ON)', 3)
    add_bullet_point(doc, "Personalized Responses: AI responds in personalized tone")
    add_bullet_point(doc, "Business Context: Uses business-specific tone templates")
    add_bullet_point(doc, "Direct Answers: Provides direct, actionable answers")
    add_bullet_point(doc, "Conversational: Natural, human-like responses")
    
    doc.add_heading('11.3.2 Categorization Mode (OFF)', 3)
    add_bullet_point(doc, "Categorization Only: AI categorizes and summarizes")
    add_bullet_point(doc, "No Personalization: Neutral, factual responses")
    add_bullet_point(doc, "Structured Output: Organized summaries and categories")
    add_bullet_point(doc, "Analysis Focus: Emphasis on analysis over conversation")
    
    doc.add_heading('11.4 Multi-Model Support', 2)
    add_bullet_point(doc, "OpenAI GPT-4: Primary model for sophisticated reasoning")
    add_bullet_point(doc, "Anthropic Claude: Alternative for different reasoning styles")
    add_bullet_point(doc, "Ollama (Local): Self-hosted models for privacy")
    add_bullet_point(doc, "Model Selection: User can choose model per GPT")
    add_bullet_point(doc, "Fallback: Automatic fallback if primary model fails")
    
    doc.add_heading('11.5 Business-Specific Features', 2)
    add_bullet_point(doc, "Fee Rate Analysis: Understand and calculate from fee spreadsheets")
    add_bullet_point(doc, "Contract Analysis: Extract and summarize contract terms")
    add_bullet_point(doc, "Policy Understanding: Answer questions about company policies")
    add_bullet_point(doc, "Financial Analysis: Analyze financial documents and spreadsheets")
    add_bullet_point(doc, "Operational Knowledge: Answer questions about business operations")
    
    doc.add_page_break()
    
    # 12. IMPLEMENTATION ROADMAP
    doc.add_heading('12. Implementation Roadmap', 1)
    
    doc.add_heading('12.1 Phase 1: Core Foundation', 2)
    add_numbered_point(doc, "Set up project structure (backend + frontend)")
    add_numbered_point(doc, "Implement authentication system (email/password + OAuth)")
    add_numbered_point(doc, "Create database schema (businesses, GPTs, conversations, documents)")
    add_numbered_point(doc, "Implement vector database integration (ChromaDB)")
    add_numbered_point(doc, "Set up document processing pipeline")
    add_numbered_point(doc, "Create basic UI layout (sidebar + main chat area)")
    
    doc.add_heading('12.2 Phase 2: GPT Management', 2)
    add_numbered_point(doc, "Implement GPT creation and storage")
    add_numbered_point(doc, "Build Edit GPT panel with all fields")
    add_numbered_point(doc, "Implement GPT kebab menu with all options")
    add_numbered_point(doc, "Create GPT vector namespace isolation")
    add_numbered_point(doc, "Implement GPT persistence (survives conversation deletion)")
    add_numbered_point(doc, "Add 'New chat' functionality (reuses GPT config)")
    
    doc.add_heading('12.3 Phase 3: Document Management', 2)
    add_numbered_point(doc, "Implement file upload (drag-and-drop + browse)")
    add_numbered_point(doc, "Add document processing (chunking + embedding)")
    add_numbered_point(doc, "Create document-GPT association system")
    add_numbered_point(doc, "Build document management UI (list, delete, preview)")
    add_numbered_point(doc, "Implement document persistence (survives conversation deletion)")
    add_numbered_point(doc, "Add support for all file formats (PDF, DOCX, XLSX, etc.)")
    
    doc.add_heading('12.4 Phase 4: Conversation Management', 2)
    add_numbered_point(doc, "Implement conversation creation and storage")
    add_numbered_point(doc, "Build conversation list in sidebar")
    add_numbered_point(doc, "Add conversation menu (rename, archive, delete)")
    add_numbered_point(doc, "Implement chat history display")
    add_numbered_point(doc, "Create message storage and retrieval")
    add_numbered_point(doc, "Ensure conversation deletion does NOT delete GPT")
    
    doc.add_heading('12.5 Phase 5: RAG & AI Integration', 2)
    add_numbered_point(doc, "Implement RAG service (retrieval + generation)")
    add_numbered_point(doc, "Integrate LLM providers (OpenAI, Anthropic, Ollama)")
    add_numbered_point(doc, "Build query processing pipeline")
    add_numbered_point(doc, "Implement source citation system")
    add_numbered_point(doc, "Add 'Reply as me' toggle functionality")
    add_numbered_point(doc, "Create tone template system")
    add_numbered_point(doc, "Implement sophisticated reasoning (not just keywords)")
    
    doc.add_heading('12.6 Phase 6: UI/UX Polish', 2)
    add_numbered_point(doc, "Fix avatar button position (bottom-left of sidebar)")
    add_numbered_point(doc, "Position '+' button next to prompt field")
    add_numbered_point(doc, "Implement file uploader above prompt (not in chat)")
    add_numbered_point(doc, "Add 'Reply as me' toggle to sidebar")
    add_numbered_point(doc, "Style all components to match ChatGPT")
    add_numbered_point(doc, "Add loading states and error handling")
    add_numbered_point(doc, "Implement responsive design")
    
    doc.add_heading('12.7 Phase 7: Settings Panel', 2)
    add_numbered_point(doc, "Create tabbed settings layout")
    add_numbered_point(doc, "Implement all 7 settings tabs")
    add_numbered_point(doc, "Add lazy loading for tabs")
    add_numbered_point(doc, "Implement settings persistence")
    add_numbered_point(doc, "Add 'Back to Chat' navigation")
    add_numbered_point(doc, "Wire all settings to backend")
    
    doc.add_heading('12.8 Phase 8: Integrations', 2)
    add_numbered_point(doc, "Build Outlook add-in (manifest + UI)")
    add_numbered_point(doc, "Create Gmail extension")
    add_numbered_point(doc, "Implement email analysis features")
    add_numbered_point(doc, "Add attachment processing in email clients")
    add_numbered_point(doc, "Test all integrations thoroughly")
    
    doc.add_heading('12.9 Phase 9: Testing & Quality Assurance', 2)
    add_numbered_point(doc, "Unit tests for all backend endpoints")
    add_numbered_point(doc, "Integration tests for RAG flow")
    add_numbered_point(doc, "UI/UX testing (10+ test cycles)")
    add_numbered_point(doc, "Security testing and penetration testing")
    add_numbered_point(doc, "Performance testing (load, stress)")
    add_numbered_point(doc, "Cross-browser testing")
    add_numbered_point(doc, "Mobile responsiveness testing")
    
    doc.add_heading('12.10 Phase 10: Deployment & Launch', 2)
    add_numbered_point(doc, "Set up production infrastructure")
    add_numbered_point(doc, "Configure CI/CD pipeline")
    add_numbered_point(doc, "Deploy backend API")
    add_numbered_point(doc, "Deploy frontend web app")
    add_numbered_point(doc, "Submit Outlook add-in to Microsoft Store")
    add_numbered_point(doc, "Submit Gmail extension to Chrome Web Store")
    add_numbered_point(doc, "Launch and monitor")
    
    doc.add_page_break()
    
    # 13. TECHNICAL REQUIREMENTS
    doc.add_heading('13. Technical Requirements', 1)
    
    doc.add_heading('13.1 Backend Requirements', 2)
    add_bullet_point(doc, "Python 3.9+")
    add_bullet_point(doc, "FastAPI framework")
    add_bullet_point(doc, "ChromaDB for vector storage")
    add_bullet_point(doc, "Unstructured.io or PyPDF2 for document processing")
    add_bullet_point(doc, "OpenAI/Anthropic API clients")
    add_bullet_point(doc, "JWT for authentication")
    add_bullet_point(doc, "PostgreSQL (for production) or JSON files (for prototype)")
    
    doc.add_heading('13.2 Frontend Requirements', 2)
    add_bullet_point(doc, "Python 3.9+")
    add_bullet_point(doc, "Streamlit framework")
    add_bullet_point(doc, "Custom CSS for styling")
    add_bullet_point(doc, "JavaScript for interactive components (if needed)")
    
    doc.add_heading('13.3 Infrastructure Requirements', 2)
    add_bullet_point(doc, "Web server (Nginx or similar)")
    add_bullet_point(doc, "Application server (Gunicorn or Uvicorn)")
    add_bullet_point(doc, "File storage (local or cloud)")
    add_bullet_point(doc, "Database server (PostgreSQL for production)")
    add_bullet_point(doc, "Vector database server (ChromaDB)")
    add_bullet_point(doc, "CDN for static assets")
    
    doc.add_heading('13.4 Security Requirements', 2)
    add_bullet_point(doc, "HTTPS/SSL certificates")
    add_bullet_point(doc, "API key authentication")
    add_bullet_point(doc, "JWT token security")
    add_bullet_point(doc, "Data encryption at rest")
    add_bullet_point(doc, "Data encryption in transit")
    add_bullet_point(doc, "Regular security audits")
    
    doc.add_heading('13.5 Performance Requirements', 2)
    add_bullet_point(doc, "Response time: < 2 seconds for chat responses")
    add_bullet_point(doc, "Document processing: < 30 seconds for typical documents")
    add_bullet_point(doc, "Concurrent users: Support 1000+ concurrent users")
    add_bullet_point(doc, "Scalability: Horizontal scaling capability")
    add_bullet_point(doc, "Caching: Implement caching for frequently accessed data")
    
    doc.add_page_break()
    
    # APPENDIX: Current State & Completed Changes
    doc.add_heading('Appendix A: Current State & Completed Changes', 1)
    
    doc.add_heading('A.1 Completed UI Components', 2)
    add_bullet_point(doc, "✅ Avatar at bottom-left of sidebar (styled as circular button)")
    add_bullet_point(doc, "✅ '+' button positioned to left of prompt field")
    add_bullet_point(doc, "✅ File uploader appears above prompt (not in chat)")
    add_bullet_point(doc, "✅ GPT kebab menu (⋮) with options: New chat, About, Edit GPT, Hide")
    add_bullet_point(doc, "✅ Conversation menu (⋯) with options: Rename, Archive, Delete")
    add_bullet_point(doc, "✅ 'Reply as me' toggle in sidebar")
    add_bullet_point(doc, "✅ Settings panel with tabbed layout")
    
    doc.add_heading('A.2 Completed Backend Features', 2)
    add_bullet_point(doc, "✅ Authentication endpoints (/api/v1/auth/login, /signup)")
    add_bullet_point(doc, "✅ Business/GPT management endpoints")
    add_bullet_point(doc, "✅ Document upload and processing endpoints")
    add_bullet_point(doc, "✅ Conversation management endpoints")
    add_bullet_point(doc, "✅ RAG service with vector database integration")
    add_bullet_point(doc, "✅ Document processor with multiple file format support")
    
    doc.add_heading('A.3 Known Issues to Fix', 2)
    add_bullet_point(doc, "❌ Login functionality needs proper error handling")
    add_bullet_point(doc, "❌ Google OAuth shows 'coming soon' - needs real implementation")
    add_bullet_point(doc, "❌ Avatar icon shows '?' instead of proper icon when logged out")
    add_bullet_point(doc, "❌ '+' button positioning needs refinement")
    add_bullet_point(doc, "❌ File uploader visibility needs better control")
    add_bullet_point(doc, "❌ GPT-Conversation lifecycle separation needs verification")
    
    doc.add_heading('A.4 Next Steps', 2)
    add_numbered_point(doc, "Fix all known issues listed above")
    add_numbered_point(doc, "Implement real OAuth flows (Google, Microsoft)")
    add_numbered_point(doc, "Complete Edit GPT panel with real backend wiring")
    add_numbered_point(doc, "Implement sophisticated reasoning (not just keyword matching)")
    add_numbered_point(doc, "Add Outlook add-in")
    add_numbered_point(doc, "Add Gmail extension")
    add_numbered_point(doc, "Comprehensive testing (10+ cycles)")
    add_numbered_point(doc, "Performance optimization")
    add_numbered_point(doc, "Security hardening")
    add_numbered_point(doc, "Production deployment")
    
    # Save document
    doc.save('AI_CoWorker_Platform_Plan.docx')
    print("Document created successfully: AI_CoWorker_Platform_Plan.docx")

if __name__ == "__main__":
    try:
        create_plan_document()
    except ImportError:
        print("python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call(["pip", "install", "python-docx"])
        create_plan_document()

